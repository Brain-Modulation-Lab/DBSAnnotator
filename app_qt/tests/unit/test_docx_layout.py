"""Tests for Word layout helpers."""

from __future__ import annotations

from docx import Document

from dbs_annotator.utils.docx_layout import (
    keep_paragraphs_with_following_block,
    keep_table_rows_together,
    keep_with_next,
)


def test_keep_with_next_sets_paragraph_property() -> None:
    doc = Document()
    paragraph = doc.add_paragraph("intro")
    keep_with_next(paragraph)
    assert paragraph.paragraph_format.keep_with_next is True


def test_keep_paragraphs_with_following_block() -> None:
    doc = Document()
    first = doc.add_paragraph("heading")
    second = doc.add_paragraph("details")
    keep_paragraphs_with_following_block([first, second])
    assert first.paragraph_format.keep_with_next is True
    assert second.paragraph_format.keep_with_next is True


def test_keep_table_rows_together_adds_cant_split() -> None:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    keep_table_rows_together(table)
    for row in table.rows:
        assert row._tr.trPr is not None
        assert row._tr.trPr.xpath(".//w:cantSplit")  # noqa: S101
