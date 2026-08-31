"""Broad SessionExporter helper and branch coverage (mocked I/O)."""

from __future__ import annotations

from unittest.mock import MagicMock, patch

import pandas as pd
import pytest
from docx import Document

from dbs_annotator.models.session_data import SessionData
from dbs_annotator.utils.session_exporter import SessionExporter


@pytest.fixture
def exporter():
    sd = MagicMock(spec=SessionData)
    sd.file_path = ""
    sd.is_file_open = MagicMock(return_value=False)
    return SessionExporter(sd), sd


def test_normalize_block_id_column_variants(exporter):
    ex, _ = exporter
    df = pd.DataFrame({"block_ID": [1, 2]})
    out = ex._normalize_block_id_column(df)
    assert "block_ID" in out.columns

    df2 = pd.DataFrame({"block_id": [1]})
    assert "block_ID" in ex._normalize_block_id_column(df2).columns

    df3 = pd.DataFrame({"blockId": [1]})
    assert "block_ID" in ex._normalize_block_id_column(df3).columns

    empty = pd.DataFrame()
    assert ex._normalize_block_id_column(empty).empty


def test_get_manufacturer_for_model(exporter):
    ex, _ = exporter
    assert ex._get_manufacturer_for_model("") == ""
    # first model in MANUFACTURERS dict
    from dbs_annotator.config_electrode_models import ELECTRODE_MODELS

    name = next(iter(ELECTRODE_MODELS))
    m = ex._get_manufacturer_for_model(name)
    assert isinstance(m, str)


def test_pick_latest_row(exporter):
    ex, _ = exporter
    assert ex._pick_latest_row(pd.DataFrame()) is None
    df = pd.DataFrame({"block_ID": [1, 3, 2], "x": [1, 2, 3]})
    row = ex._pick_latest_row(df)
    assert int(row["block_ID"]) == 3


def test_pick_latest_session_row(exporter):
    ex, _ = exporter
    assert ex._pick_latest_session_row(pd.DataFrame()) is None
    df = pd.DataFrame(
        {
            "session_ID": [1, 2, 2],
            "block_ID": [1, 1, 2],
            "scale_value": [1, 2, 3],
        }
    )
    r = ex._pick_latest_session_row(df)
    assert r is not None


def test_column_header(exporter):
    ex, _ = exporter
    assert isinstance(ex._column_header("scale_name"), str)


def test_read_session_data_none_when_no_path(exporter):
    ex, sd = exporter
    sd.file_path = None
    assert ex._read_session_data() is None


def test_read_session_data_reads_tsv(tmp_path, exporter):
    ex, sd = exporter
    p = tmp_path / "d.tsv"
    p.write_text("a\tb\n1\t2\n", encoding="utf-8")
    sd.file_path = str(p)
    df = ex._read_session_data()
    assert df is not None
    assert len(df) == 1


def test_add_summary_section_with_notes(exporter):
    ex, _ = exporter
    doc = Document()
    df = pd.DataFrame()
    df_init = pd.DataFrame(
        {
            "session_ID": [1],
            "scale_name": ["Y"],
            "scale_value": ["1"],
            "notes": ["hello"],
            "block_ID": [0],
        }
    )
    ex._add_summary_section(doc, df, df_init, df)
    assert len(doc.paragraphs) >= 1


def test_add_programming_summary_empty_df(exporter):
    ex, _ = exporter
    doc = Document()
    ex._add_programming_summary(doc, pd.DataFrame(), pd.DataFrame(), pd.DataFrame())
    assert any("No session" in p.text for p in doc.paragraphs)


def test_add_programming_summary_parses_split_amplitude_and_numeric_text(exporter):
    ex, _ = exporter
    doc = Document()
    df = pd.DataFrame(
        {
            "block_ID": [1, 2],
            "left_amplitude": ["2.5_1.5", "3.0_2.0"],
            "right_amplitude": ["1.0_1.0", "1.5_1.5"],
            "left_stim_freq": ["130", "140 Hz"],
            "right_stim_freq": ["130", "150"],
            "left_pulse_width": ["60 µs", "70"],
            "right_pulse_width": ["80", "90 µs"],
        }
    )

    ex._add_programming_summary(doc, df, pd.DataFrame(), pd.DataFrame())
    summary_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Amplitude Range:  L: 4.0 - 5.0 mA  |  R: 2.0 - 3.0 mA" in summary_text
    assert "Frequency Range:  L: 130 - 140 Hz  |  R: 130 - 150 Hz" in summary_text
    assert "Pulse Width Range:  L: 60 - 70 µs  |  R: 80 - 90 µs" in summary_text


def test_find_best_and_second_best_blocks_empty(exporter):
    ex, _ = exporter
    assert ex._find_best_and_second_best_blocks(pd.DataFrame()) == ([], [])


def test_find_best_and_second_best_blocks_minimal(exporter):
    ex, _ = exporter
    ex.set_scale_optimization_prefs([("Mood", "0", "10", "max", "")])
    df = pd.DataFrame(
        {
            "block_ID": [1, 1],
            "scale_name": ["Mood", "Mood"],
            "scale_value": ["5", "8"],
            "laterality": ["L", "L"],
        }
    )
    a, b = ex._find_best_and_second_best_blocks(df)
    assert isinstance(a, list)
    assert isinstance(b, list)


def test_export_annotations_to_word_cancel_dialog(tmp_path, exporter):
    ex, sd = exporter
    sd.is_file_open.return_value = True
    sd.file_path = str(tmp_path / "a.tsv")
    with patch(
        "PySide6.QtWidgets.QFileDialog.getSaveFileName",
        return_value=("", ""),
    ):
        assert ex.export_annotations_to_word() is False


def test_export_annotations_to_pdf_cancel_dialog(tmp_path, exporter):
    ex, sd = exporter
    sd.is_file_open.return_value = True
    sd.file_path = str(tmp_path / "a.tsv")
    with patch(
        "PySide6.QtWidgets.QFileDialog.getSaveFileName",
        return_value=("", ""),
    ):
        assert ex.export_annotations_to_pdf() is False


def test_export_longitudinal_report_controller_calls_exporter(monkeypatch):
    from dbs_annotator.controllers.wizard_controller import WizardController

    c = WizardController()
    called = {}

    class FakeExp:
        def set_scale_optimization_prefs(self, p):
            called["prefs"] = p

        def set_clinical_scale_prefs(self, p):
            called["clinical"] = p

        def export_to_word(self, *a, **k):
            called["word"] = True

        def export_to_pdf(self, *a, **k):
            called["pdf"] = True

    def fake_exporter():
        return FakeExp()

    monkeypatch.setattr(
        "dbs_annotator.utils.longitudinal_exporter.LongitudinalExporter",
        fake_exporter,
    )
    c.export_longitudinal_report([], [], "word")
    assert called.get("word")
    c.export_longitudinal_report([], [], "pdf")
    assert called.get("pdf")


def _lateral_row(**overrides):
    """One session TSV row with distinct left/right stimulation values."""
    row = {
        "date": "2026-08-20",
        "time": "10:00:00",
        "block_ID": "1",
        "session_ID": "1",
        "is_initial": "0",
        "scale_name": "Tremor",
        "scale_value": "3",
        "electrode_model": "Medtronic 3389",
        "program_ID": "A",
        "left_stim_freq": "130",
        "left_anode": "case",
        "left_cathode": "E1",
        "left_amplitude": "2.0",
        "left_pulse_width": "60",
        "right_stim_freq": "125",
        "right_anode": "case",
        "right_cathode": "E2",
        "right_amplitude": "1.5",
        "right_pulse_width": "90",
        "notes": "ok",
    }
    row.update(overrides)
    return row


LATERAL_COLUMNS = ("frequency", "anode", "cathode", "amplitude", "pulse_width")


def test_create_lateral_table_data_keeps_both_sides(exporter):
    """Regression: the R row must carry every lateral parameter, not just
    pulse_width.

    Two separate loops used to split the left and right columns, and the second
    reused a stale ``generic_col`` from the first. Every right-hand value
    therefore overwrote ``pulse_width`` and the R row rendered blank frequency,
    anode, cathode and amplitude in the report table.
    """
    ex, _ = exporter
    out = ex._create_lateral_table_data(pd.DataFrame([_lateral_row()]))

    assert list(out["laterality"]) == ["L", "R"]
    left = out[out["laterality"] == "L"].iloc[0]
    right = out[out["laterality"] == "R"].iloc[0]

    for col in LATERAL_COLUMNS:
        assert pd.notna(left[col]) and left[col] != "", f"L row lost {col}"
        assert pd.notna(right[col]) and right[col] != "", f"R row lost {col}"

    assert left["frequency"] == "130"
    assert left["cathode"] == "E1"
    assert left["amplitude"] == "2.0"
    assert left["pulse_width"] == "60"

    assert right["frequency"] == "125"
    assert right["cathode"] == "E2"
    assert right["amplitude"] == "1.5"
    assert right["pulse_width"] == "90"


def test_create_lateral_table_data_groups_scales_per_block(exporter):
    """Two scales in one block collapse into newline-joined common cells, and
    the block still yields exactly one L/R pair."""
    ex, _ = exporter
    df = pd.DataFrame(
        [
            _lateral_row(scale_name="Tremor", scale_value="3"),
            _lateral_row(scale_name="Rigidity", scale_value="5"),
        ]
    )
    out = ex._create_lateral_table_data(df)

    assert len(out) == 2
    assert list(out["laterality"]) == ["L", "R"]
    for _, row in out.iterrows():
        assert row["scale_name"] == "Tremor\nRigidity"
        assert row["scale_value"] == "3\n5"
