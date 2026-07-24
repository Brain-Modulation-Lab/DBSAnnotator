"""Shared helpers for the session and longitudinal report exporters.

These were duplicated (near byte-for-byte) across ``session_exporter`` and
``longitudinal_exporter``; they live here once and both exporters delegate to
them.  Everything here is a free function (no exporter state) except where a
value is passed in explicitly.
"""

from __future__ import annotations

import importlib
import os
import shutil
import subprocess
import sys
import tempfile
from typing import Protocol, cast

from docx.document import Document as DocumentType
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from PySide6.QtCore import Qt, QTimer
from PySide6.QtGui import QColor, QPainter, QPixmap
from PySide6.QtWidgets import QMessageBox, QWidget

from ..config_electrode_models import ELECTRODE_MODELS, MANUFACTURERS, ContactState
from ..models import ElectrodeCanvas


class _ExportTransientParent(Protocol):
    """QWidget host that may store a reference to a transient export message."""

    _export_transient_msg: QMessageBox | None


def show_transient_message(
    parent: QWidget | None,
    title: str,
    text: str,
    *,
    msecs: int = 2000,
    icon: QMessageBox.Icon = QMessageBox.Icon.Information,
) -> None:
    """Show a non-modal message box that auto-closes after *msecs*."""
    msg = QMessageBox(parent)
    msg.setIcon(icon)
    msg.setWindowTitle(title)
    msg.setText(text)
    msg.setStandardButtons(QMessageBox.StandardButton.NoButton)
    msg.setWindowModality(Qt.WindowModality.NonModal)
    msg.show()

    if parent is not None:
        try:
            cast(_ExportTransientParent, parent)._export_transient_msg = msg
        except Exception:
            pass

    # Use a dedicated QTimer owned by the message box so it reliably fires.
    # Some Qt builds do not ship QWeakPointer.
    timer = QTimer(msg)
    timer.setSingleShot(True)

    def _close_msg() -> None:
        try:
            msg.accept()
        except Exception:
            try:
                msg.close()
            except Exception:
                pass

        if parent is not None:
            try:
                host = cast(_ExportTransientParent, parent)
                if host._export_transient_msg is msg:
                    host._export_transient_msg = None
            except Exception:
                pass

    timer.timeout.connect(_close_msg)
    timer.start(max(0, int(msecs)))


def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> None:
    """Convert a Word document to PDF using the best available method.

    Tries in order:
    1. docx2pdf (requires Microsoft Word COM)
    2. Word COM via PowerShell subprocess
    3. LibreOffice headless

    Raises RuntimeError if no conversion method succeeds.
    """
    errors: list[str] = []

    # 1. Try docx2pdf. Imported dynamically: it is an optional dependency only
    # installed on Windows/macOS (see pyproject), so a static import would fail
    # type-checking on Linux where the package is intentionally absent.
    try:
        convert = importlib.import_module("docx2pdf").convert

        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            return
    except Exception as exc:
        errors.append(f"docx2pdf: {exc}")

    # 2. Try Word COM via PowerShell (Windows)
    try:
        abs_docx = os.path.abspath(docx_path).replace("'", "''")
        abs_pdf = os.path.abspath(pdf_path).replace("'", "''")
        ps_script = (
            "$w = New-Object -ComObject Word.Application; "
            "$w.Visible = $false; "
            f"$d = $w.Documents.Open('{abs_docx}'); "
            f"$d.SaveAs2('{abs_pdf}', 17); "
            "$d.Close(); $w.Quit()"
        )
        subprocess.run(
            ["powershell", "-NoProfile", "-Command", ps_script],
            check=True,
            capture_output=True,
            timeout=60,
        )
        if os.path.exists(pdf_path):
            return
    except Exception as exc:
        errors.append(f"Word COM (PowerShell): {exc}")

    # 3. Try LibreOffice headless
    soffice = shutil.which("soffice")
    if soffice:
        try:
            out_dir = os.path.dirname(os.path.abspath(pdf_path))
            subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    out_dir,
                    os.path.abspath(docx_path),
                ],
                check=True,
                capture_output=True,
                timeout=60,
            )
            # LibreOffice outputs with same basename
            lo_output = os.path.join(
                out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf"
            )
            if lo_output != pdf_path and os.path.exists(lo_output):
                shutil.move(lo_output, pdf_path)
            if os.path.exists(pdf_path):
                return
        except Exception as exc:
            errors.append(f"LibreOffice: {exc}")
    else:
        errors.append("LibreOffice: soffice not found on PATH")

    detail = "\n".join(errors)
    raise RuntimeError(
        f"Could not convert to PDF. Tried all available methods:\n{detail}\n\n"
        "Please export to Word (.docx) and convert to PDF manually."
    )


def open_file(path: str) -> None:
    """Open a file with the system default application."""
    try:
        if sys.platform == "win32":
            os.startfile(path)  # noqa: S606
        elif sys.platform == "darwin":
            subprocess.Popen(["open", path])  # noqa: S603
        else:
            subprocess.Popen(["xdg-open", path])  # noqa: S603
    except Exception:
        pass


def set_cell_border_top(cell, sz: int = 12) -> None:
    """Set the top border of a table cell (sz in eighths of a point)."""
    try:
        tcPr = cell._tc.get_or_add_tcPr()  # noqa: N806
        tcBorders = OxmlElement("w:tcBorders")  # noqa: N806
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), str(sz))
        top.set(qn("w:space"), "0")
        top.set(qn("w:color"), "000000")
        tcBorders.append(top)
        tcPr.append(tcBorders)
    except Exception:
        pass


def highlight_cells(row_cells, intensity: str = "best") -> None:
    """Apply a green background to all cells in a row.

    ``intensity`` is "best" (darker green) or "second" (lighter green).
    """
    color = "96D2A0" if intensity == "best" else "C8EBCD"
    for cell in row_cells:
        try:
            shading_elm = OxmlElement("w:shd")
            shading_elm.set(qn("w:fill"), color)
            cell._tc.get_or_add_tcPr().append(shading_elm)
        except Exception:
            pass


def get_manufacturer_for_model(model_name: str) -> str:
    """Return the manufacturer string for a given electrode model name."""
    if not model_name:
        return ""
    for manufacturer, models in (MANUFACTURERS or {}).items():
        try:
            if model_name in models:
                return str(manufacturer)
        except Exception:
            continue
    return ""


def add_table_legend(
    doc: DocumentType,
    best_ids: list,
    second_ids: list,
    scale_optimization_prefs: list,
    *,
    entry_noun: str,
) -> None:
    """Add the colour legend, scale-target summary, and clinical disclaimer.

    ``entry_noun`` is the wording used in the legend ("configuration" for the
    session report, "entry" for the longitudinal report).
    """
    if not best_ids and not second_ids:
        return

    doc.add_paragraph()  # spacing

    legend_para = doc.add_paragraph()
    legend_para.add_run("Legend: ").bold = True

    if best_ids:
        best_run = legend_para.add_run("■ ")
        best_run.font.color.rgb = RGBColor(0x96, 0xD2, 0xA0)
        legend_para.add_run(f"Optimal {entry_noun}    ")

    if second_ids:
        second_run = legend_para.add_run("■ ")
        second_run.font.color.rgb = RGBColor(0xC8, 0xEB, 0xCD)
        legend_para.add_run(f"Second-best {entry_noun}")

    # Show target values used for optimization
    if scale_optimization_prefs:
        targets_para = doc.add_paragraph()
        targets_para.add_run("Scale targets: ").bold = True
        target_parts = []
        for pref in scale_optimization_prefs:
            if len(pref) >= 5:
                name, smin, smax, mode, custom_val = pref
                if mode == "ignore":
                    continue
                elif mode == "min":
                    target_parts.append(f"{name}: min")
                elif mode == "max":
                    target_parts.append(f"{name}: max")
                elif mode == "custom":
                    target_parts.append(f"{name}: {custom_val}")
        if target_parts:
            targets_para.add_run("; ".join(target_parts))
            for run in targets_para.runs:
                run.font.size = Pt(9)

    # Clinical disclaimer
    disclaimer_para = doc.add_paragraph()
    disclaimer_run = disclaimer_para.add_run(
        "Note: The highlighted rows are derived exclusively from the recorded "
        "session scale values and represent a computational ranking intended "
        "solely as a reference. This color-coded indication does not constitute "
        "clinical guidance."
    )
    disclaimer_run.font.size = Pt(9)
    disclaimer_run.font.italic = True


def render_electrode_png(
    model_name: str,
    anode_text: str,
    cathode_text: str,
    target_size_px: tuple[int, int] = (440, 900),
) -> str | None:
    """Render an electrode configuration to a temporary white-background PNG.

    Returns the temp file path, or None if the model is unknown / rendering
    fails.
    """
    try:
        model = ELECTRODE_MODELS.get(model_name)
        if not model:
            return None

        canvas = ElectrodeCanvas()
        canvas.set_model(model)
        canvas.resize(*target_size_px)
        try:
            canvas.set_export_mode(True)
        except Exception:
            pass

        # Apply contact states
        canvas.contact_states.clear()
        canvas.case_state = ContactState.OFF

        def apply_tokens(text: str, state: int) -> None:
            if not text:
                return
            for token in str(text).split("_"):
                token = token.strip()
                if not token:
                    continue
                if token == "case":
                    canvas.case_state = state
                    continue
                if token.startswith("E") and len(token) >= 2:
                    try:
                        if token[-1].isalpha():
                            idx = int(token[1:-1])
                            seg_map = {"a": 0, "b": 1, "c": 2}
                            seg_char = token[-1].lower()
                            if seg_char in seg_map:
                                canvas.contact_states[(idx, seg_map[seg_char])] = state
                        else:
                            idx = int(token[1:])
                            if model.is_directional:
                                for seg in range(3):
                                    canvas.contact_states[(idx, seg)] = state
                            else:
                                canvas.contact_states[(idx, 0)] = state
                    except Exception:
                        continue

        apply_tokens(anode_text, ContactState.ANODIC)
        apply_tokens(cathode_text, ContactState.CATHODIC)
        canvas.update()

        # Render with white background
        original_paint = canvas.paintEvent

        def white_bg_paint(event):
            painter = QPainter(canvas)
            painter.fillRect(canvas.rect(), Qt.GlobalColor.white)
            original_paint(event)

        canvas.paintEvent = white_bg_paint  # type: ignore[assignment]  # ty: ignore[invalid-assignment]

        pixmap = QPixmap(canvas.size())
        pixmap.fill(Qt.GlobalColor.white)
        canvas.render(pixmap)

        # Crop white borders
        image = pixmap.toImage()
        white_rgb = QColor(Qt.GlobalColor.white).rgb()
        left, top, right, bottom = image.width(), image.height(), 0, 0
        for y in range(image.height()):
            for x in range(image.width()):
                if image.pixel(x, y) != white_rgb:
                    left = min(left, x)
                    top = min(top, y)
                    right = max(right, x)
                    bottom = max(bottom, y)
        if right > left and bottom > top:
            margin = 20
            left = max(0, left - margin)
            top = max(0, top - margin)
            right = min(image.width() - 1, right + margin)
            bottom = min(image.height() - 1, bottom + margin)
            cropped = pixmap.copy(left, top, right - left + 1, bottom - top + 1)
        else:
            cropped = pixmap

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        tmp.close()
        cropped.save(tmp.name, "PNG")
        return tmp.name
    except Exception:
        return None
